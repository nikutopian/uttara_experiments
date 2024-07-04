import os
import mimetypes
import zipfile
import PyPDF2
from PyPDF2 import PdfReader
import pytesseract
from PIL import Image
import docx
import openpyxl
import io
import logging
import time
from functools import wraps
from tqdm import tqdm
import pdf2image
import platform

if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

# Try to import magic, but provide a fallback if it's not available
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logging.warning("python-magic is not installed. Falling back to mimetypes for file type detection.")

MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB
MAX_PROCESSING_TIME = 300  # 5 minutes

class FileParserError(Exception):
    """Base exception class for FileParser errors"""
    pass

class FileTooLargeError(FileParserError):
    """Raised when the file is too large to process"""
    pass

class ProcessingTimeoutError(FileParserError):
    """Raised when processing takes too long"""
    pass

class UnsupportedFileTypeError(FileParserError):
    """Raised when an unsupported file type is encountered"""
    pass

class ParseError(FileParserError):
    """Raised when there's an error parsing a file"""
    pass

def retry(max_attempts=3, delay=1):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            attempts = 0
            while attempts < max_attempts:
                try:
                    return func(*args, **kwargs)
                except (IOError, ConnectionError) as e:
                    attempts += 1
                    if attempts == max_attempts:
                        raise
                    time.sleep(delay)
                    logging.warning(f"Retrying {func.__name__} (attempt {attempts}/{max_attempts})")
        return wrapper
    return decorator

class FileParser:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        if MAGIC_AVAILABLE:
            self.mime = magic.Magic(mime=True)
        else:
            self.mime = None

    def check_file_size(self, file_path):
        """Check if the file size is within the allowed limit."""
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            raise FileTooLargeError(f"File is too large ({file_size} bytes). Maximum allowed size is {MAX_FILE_SIZE} bytes.")

    @retry(max_attempts=3, delay=2)
    def identify_format(self, file_path):
        """Identify file format based on extension and content."""
        try:
            extension_guess = mimetypes.guess_type(file_path)[0]
            
            if MAGIC_AVAILABLE:
                content_guess = self.mime.from_file(file_path)
            else:
                # Fallback to using the file extension if magic is not available
                content_guess = extension_guess
            
            return {
                "extension_guess": extension_guess,
                "content_guess": content_guess
            }
        except Exception as e:
            self.logger.error(f"Error identifying file format: {str(e)}")
            raise FileParserError(f"Failed to identify file format: {str(e)}")

    def parse_file(self, file_path):
        """Parse file based on its format."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        self.check_file_size(file_path)

        try:
            start_time = time.time()
            file_format = self.identify_format(file_path)
            
            parsing_func = None
            if "zip" in file_format["content_guess"]:
                parsing_func = self.parse_zip
            elif "pdf" in file_format["content_guess"]:
                parsing_func = self.parse_pdf
            elif "text" in file_format["content_guess"]:
                parsing_func = self.parse_text
            elif "word" in file_format["content_guess"] or file_path.endswith('.docx'):
                parsing_func = self.parse_docx
            elif "excel" in file_format["content_guess"] or file_path.endswith('.xlsx'):
                parsing_func = self.parse_xlsx
            elif "image" in file_format["content_guess"]:
                parsing_func = self.parse_image
            else:
                raise UnsupportedFileTypeError(f"Unsupported file format: {file_format}, File: {file_path}")

            with tqdm(total=100, desc=f"Parsing file: {file_path}", unit="%") as pbar:
                result = ""
                for chunk in parsing_func(file_path):
                    result += chunk
                    pbar.update(10)  # Update progress bar
                    if time.time() - start_time > MAX_PROCESSING_TIME:
                        raise ProcessingTimeoutError(f"Processing time exceeded {MAX_PROCESSING_TIME} seconds")
                pbar.update(100 - pbar.n)  # Ensure the progress bar reaches 100%
            return result

        except FileParserError:
            raise
        except Exception as e:
            self.logger.error(f"Unexpected error parsing file: {str(e)}")
            raise ParseError(f"Failed to parse file: {str(e)}")

    @retry(max_attempts=3, delay=2)
    def parse_zip(self, file_path):
        """Parse zip file and extract its contents."""
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                extract_dir = file_path + "_extracted"
                total_files = len(zip_ref.namelist())
                for i, file in enumerate(zip_ref.namelist()):
                    zip_ref.extract(file, extract_dir)
                    yield f"Extracted {i+1}/{total_files}: {file}\n"
        except zipfile.BadZipFile:
            raise ParseError("The file is not a valid ZIP archive")
        except PermissionError:
            raise ParseError("Permission denied when trying to extract ZIP contents")
        except Exception as e:
            raise ParseError(f"Error extracting ZIP file: {str(e)}")

    @retry(max_attempts=3, delay=2)
    def parse_pdf(self, file_path):
        """Parse PDF file, including OCR for images."""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                total_pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text() or ''
                    if not text.strip():  # If no text extracted, try OCR
                        text = self.ocr_pdf_page(page)
                    yield f"Page {i+1}/{total_pages}: {text}\n"
        except PyPDF2.errors.PdfReadError:
            raise ParseError("The file is not a valid PDF or is encrypted")
        except Exception as e:
            raise ParseError(f"Error parsing PDF file: {str(e)}")


    def ocr_from_pdf(self, pdf_path):
        pages = pdf2image.convert_from_path(pdf_path)
        text = ""
        for page in pages:
            text += pytesseract.image_to_string(page)
        return text        


    def ocr_pdf_page(self, page):
        """Perform OCR on a PDF page."""
        try:
            if '/XObject' in page['/Resources']:
                xobjects = page['/Resources']['/XObject'].get_object()
                for obj in xobjects:
                    xobject = xobjects[obj]
                    if xobject['/Subtype'] == '/Image':
                        size = (xobject['/Width'], xobject['/Height'])
                        data = xobject.get_data()
                        colorspace = xobject['/ColorSpace']
                        bits_per_component = xobject['/BitsPerComponent']
                        filters = xobject.get('/Filter', [])

                        if isinstance(filters, str):
                            filters = [filters]

                        if '/DCTDecode' in filters:
                            # This is a JPEG image
                            img = Image.open(io.BytesIO(data))
                        else:
                            # Determine the appropriate mode based on colorspace
                            if colorspace == '/DeviceRGB':
                                mode = 'RGB'
                            elif colorspace == '/DeviceGray':
                                mode = 'L'
                            else:
                                self.logger.warning(f"Unsupported colorspace: {colorspace}")
                                continue

                            # Handle different bits per component
                            if bits_per_component == 8:
                                img = Image.frombytes(mode, size, data)
                            elif bits_per_component == 1:
                                img = Image.frombytes('1', size, data)
                                img = img.convert('L')  # Convert to grayscale for better OCR
                            else:
                                self.logger.warning(f"Unsupported bits per component: {bits_per_component}")
                                continue

                        # Perform OCR on the image
                        return pytesseract.image_to_string(img)

            # If no image found in direct XObjects, check for Form XObjects
            for key, value in page['/Resources']['/XObject'].items():
                if isinstance(value, dict) and value.get('/Subtype') == '/Form':
                    # Extract the content stream of the Form XObject
                    content = value.get_object().get('/Contents', b'')
                    if isinstance(content, bytes):
                        # Create a new PDF page from the Form XObject content
                        temp_pdf = PdfReader(io.BytesIO(b'%PDF-1.7\n' + content + b'\n%%EOF'))
                        temp_page = temp_pdf.pages[0]
                        # Recursively call ocr_pdf_page on the new page
                        return self.ocr_pdf_page(temp_page)

            self.logger.warning("No image found in the PDF page for OCR.")
            return ""
        except Exception as e:
            self.logger.error(f"OCR failed for PDF page: {str(e)}")
            return ""    

    @retry(max_attempts=3, delay=2)
    def parse_text(self, file_path):
        """Parse text file."""
        try:
            with open(file_path, 'r', errors='ignore') as file:
                for line in file:
                    yield line
        except UnicodeDecodeError:
            raise ParseError("Unable to decode the text file. It might be in an unsupported encoding.")
        except Exception as e:
            raise ParseError(f"Error reading text file: {str(e)}")

    @retry(max_attempts=3, delay=2)
    def parse_docx(self, file_path):
        """Parse DOCX file."""
        try:
            doc = docx.Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            for i, paragraph in enumerate(doc.paragraphs):
                yield f"{paragraph.text}\n"
        except docx.opc.exceptions.PackageNotFoundError:
            raise ParseError("The file is not a valid DOCX document")
        except Exception as e:
            raise ParseError(f"Error parsing DOCX file: {str(e)}")

    @retry(max_attempts=3, delay=2)
    def parse_xlsx(self, file_path):
        """Parse XLSX file."""
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                total_rows = sheet.max_row
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    yield f"Sheet '{sheet_name}', Row {i+1}/{total_rows}: {' | '.join(str(cell) for cell in row if cell is not None)}\n"
        except openpyxl.utils.exceptions.InvalidFileException:
            raise ParseError("The file is not a valid XLSX spreadsheet")
        except Exception as e:
            raise ParseError(f"Error parsing XLSX file: {str(e)}")

    @retry(max_attempts=3, delay=2)
    def parse_image(self, file_path):
        """Parse image file using OCR."""
        try:
            img = Image.open(file_path)
            yield pytesseract.image_to_string(img)
        except Image.UnidentifiedImageError:
            raise ParseError("The file is not a valid image or the image format is not supported")
        except Exception as e:
            raise ParseError(f"Error parsing image file: {str(e)}")

